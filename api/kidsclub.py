
"""
 Program is developed by Joshua Yalung.
This program generates a sign-out document for
 students who signed in for the day at Kids Club.
 We aim for this program to operate with other software/databases that
 Saint Edward Catholic School uses at its disposal, such as Procare.
    This program is a part of the Kids Club Check-in System.

"""
import docx
import csv
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from datetime import date

SHEET_PATH = "C:/Users/jcyal/Documents/stedcskidsclub/signInSheets/"
current_date = date.today()
def get_students():
    students = []
    with open('students.csv', mode = 'r') as csv_file:
        students_file = csv.reader(csv_file)
        for student in students_file:
            if(student[0] == 'Last Name'):
                continue
            students.append([student[0], student[1], student[2]])
    return students

def save_document(name : str, students: list):
    print(students)
    # gathering today's date for sign out
    
    date_string = current_date.strftime("%A, %B %d, %Y")
    # generate document
    document = Document()

    # header for sign out sheet
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Saint Edward School Kids Club"
                            "\nSign-Out Sheet"
                            "\n" + date_string +
                            "\n # of Kids: " + str(len(students)))
    run.bold = True

    # font variable
    font = run.font

    # change font to Times New Roman and size to 12
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set header font to black
    font_color = font.color
    font_color.rgb = None

    # sort the list in alphabetical order from last name
    students.sort(key=lambda student: student[0])

    # generate table and create the header for the document
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STUDENT\'S NAME'
    hdr_cells[1].text = 'TIME OUT'
    hdr_cells[2].text = 'GUARDIAN\'S SIGNATURE'

    # variable for time character
    time_character = ':'

    # padding check
    num_entries = 2
    # iterate through the list of students and insert them into the table
    for student in students:
        temp = [student[0] + ", " + student[1]]
        # insert a row into the table and format the table
        row = table.add_row()
        row.height = Inches(0.25)
        row_cells = row.cells

        # insert student names into first column
        cell1 = row_cells[0].add_paragraph(temp)

        # time cell
        cell2 = row_cells[1].add_paragraph()
        cell2_run = cell2.add_run(time_character)
        cell2_run.font.size = Pt(24)

        # align cells to center format
        cell1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_entries += 1
        if num_entries == 15:
            # add padding 
            num_entries = 0
            temp_row = table.add_row()

    # extra cells in case students come in later
    for i in range(0,3):
        row = table.add_row()
        row.height = Inches(0.25)
        row_cells = row.cells

        # insert student names into first column
        cell1 = row_cells[0].add_paragraph()

        # time cell
        cell2 = row_cells[1].add_paragraph()
        cell2_run = cell2.add_run(time_character)
        cell2_run.font.size = Pt(24)

        # align cells to center format
        cell1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # save document into folder
    document.save(SHEET_PATH + current_date.strftime("%B") + "/" + name + ".docx")












